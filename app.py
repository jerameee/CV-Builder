from os import name
from docx import Document
from docx.shared import Inches

document = Document()

# Profile Picture
document.add_picture(
    'me.png',
    width=Inches(2.0)
)

# Name Phone Number and Email Information
name = input('What is your name?')
phone_number = input('What is your phone number?')
email = input('What is your email?')
document.add_paragraph(name + '|' + phone_number + '|' + email)

# About Me
document.add_heading('About Me')
about_me = input('Tell me about yourself.')
document.add_paragraph(about_me)

# Skills
document.add_heading('Skills')
skills = input('Enter skill')
p = document.add_paragraph(skills)
p.style = 'List Bullet'

while True:
    has_more_skills = input('Do you have more skills? Yes or No')
    if has_more_skills.lower() == 'yes':
        skills = input('Enter skill')
        p = document.add_paragraph(skills)
        p.style = 'List Bullet'
    else:
        break


# Work Experience
document.add_heading('Work Experience')
p = document.add_paragraph()

company = input('Enter company')
from_date = input('From Date')
to_date = input('To Date')

p.add_run(company + ' ').bold = True
p.add_run(from_date + '-' + to_date + '\n').italic = True

experience_details = input('Describe your work experience' + company)
p.add_run(experience_details)

# Footer
section = document.sections[0]


document.save('cv.docx')


